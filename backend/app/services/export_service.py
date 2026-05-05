import io
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

# ── Shared header data ──────────────────────────────────────────────────────

HEADER = {
    "name": "Jack Crawford",
    "location": "Denton, TX",
    "phone": "(469) 715-7415",
    "email": "jhcrawford27@gmail.com",
    "github_url": "https://github.com/jhc027",
    "github_label": "GitHub",
    "linkedin_url": "https://www.linkedin.com/in/jack-crawford-54764234a/",
    "linkedin_label": "LinkedIn",
}


def _today() -> str:
    return date.today().strftime("%B %-d, %Y")


# ── DOCX ────────────────────────────────────────────────────────────────────

def generate_docx(body_text: str, company: str | None, role: str | None) -> bytes:
    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(72)
        section.left_margin = section.right_margin = Pt(72)

    def add(text: str = "", bold: bool = False, size: int = 12, color: RGBColor | None = None, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        if color:
            run.font.color.rgb = color
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    def add_hyperlink(para_text: str, url: str, label: str) -> None:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

        # Add preceding text
        if para_text:
            run = p.add_run(para_text)
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

        # Hyperlink element
        part = doc.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        new_run.append(rPr)
        text_el = OxmlElement("w:t")
        text_el.text = label
        new_run.append(text_el)
        hyperlink.append(new_run)
        p._p.append(hyperlink)

    # Header
    add(HEADER["location"])
    add(HEADER["phone"])
    add(HEADER["email"])

    # GitHub | LinkedIn line
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _add_hyperlink_run(paragraph, url, label):
        part = doc.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        f_size = OxmlElement("w:sz")
        f_size.set(qn("w:val"), "24")
        rPr.append(f_size)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = label
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    _add_hyperlink_run(p, HEADER["github_url"], HEADER["github_label"])
    sep = p.add_run(" | ")
    sep.font.size = Pt(12)
    sep.font.name = "Times New Roman"
    _add_hyperlink_run(p, HEADER["linkedin_url"], HEADER["linkedin_label"])

    add(_today())
    add()  # blank line

    add("Dear Hiring Manager,")
    add()

    # Body paragraphs
    for para in body_text.strip().split("\n\n"):
        stripped = para.strip()
        if stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(stripped)
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

    add()
    add("Sincerely,")
    add()
    add(HEADER["name"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── PDF ─────────────────────────────────────────────────────────────────────

def generate_pdf(body_text: str, company: str | None, role: str | None) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    styles = getSampleStyleSheet()
    normal = styles["Normal"]

    name_style = ParagraphStyle("name", parent=normal, fontSize=14, leading=18, fontName="Times-Bold", spaceAfter=2)
    header_style = ParagraphStyle("header", parent=normal, fontSize=12, leading=15, fontName="Times-Roman", spaceAfter=1)
    link_style = ParagraphStyle("link", parent=normal, fontSize=12, leading=15, fontName="Times-Roman", textColor=(0.1, 0.4, 0.8), spaceAfter=1)
    date_style = ParagraphStyle("date", parent=normal, fontSize=12, leading=15, fontName="Times-Roman", spaceAfter=6)
    body_style = ParagraphStyle("body", parent=normal, fontSize=12, leading=17, fontName="Times-Roman", spaceAfter=8)
    closing_style = ParagraphStyle("closing", parent=normal, fontSize=12, leading=15, fontName="Times-Roman", spaceAfter=2)

    story = [
        Paragraph(HEADER["location"], header_style),
        Paragraph(HEADER["phone"], header_style),
        Paragraph(HEADER["email"], header_style),
        Paragraph(
            f'<a href="{HEADER["github_url"]}" color="blue">{HEADER["github_label"]}</a>'
            f' | '
            f'<a href="{HEADER["linkedin_url"]}" color="blue">{HEADER["linkedin_label"]}</a>',
            link_style,
        ),
        Paragraph(_today(), date_style),
        Spacer(1, 0.15 * inch),
        Paragraph("Dear Hiring Manager,", body_style),
    ]

    for para in body_text.strip().split("\n\n"):
        stripped = para.strip()
        if stripped:
            story.append(Paragraph(stripped, body_style))

    story += [
        Spacer(1, 0.05 * inch),
        Paragraph("Sincerely,", closing_style),
        Spacer(1, 0.25 * inch),
        Paragraph(HEADER["name"], closing_style),
    ]

    doc.build(story)
    buf.seek(0)
    return buf.read()
